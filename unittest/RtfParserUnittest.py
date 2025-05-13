import os
import yaml
import subprocess
from docx import Document
from docx.oxml import CT_P, CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from docx.oxml import OxmlElement
import enum
import re
import logging
from logging.handlers import RotatingFileHandler
import sys
from collections import defaultdict

from pandas.io.sas.sas_constants import column_name_text_subheader_length

# 配置文件路径
YAML_CONFIG = "Apnea2.yml"


class tableType(enum.Enum):
    Null = enum.auto()
    Info = enum.auto()
    FirstOrder = enum.auto()
    SleepStage = enum.auto()
    Arousal = enum.auto()
    Apnea1 = enum.auto()
    Apnea2 = enum.auto()
    LimbMovements = enum.auto()
    BreathingEvent = enum.auto()
    OxygenSaturation = enum.auto()
    Snoring = enum.auto()

class RtfDataParser:
    def __init__(self):
        self.master_columns = {}

    def judge_table_type(self,table):
        keyword_info = ["姓名"]
        keyword_firstorder = ["熄灯时间", "睡眠期平均心率"]
        # 睡眠分期
        keyword_sleepStage = ["睡眠时间"]
        # 微觉醒类型
        keyword_arousal = ["微觉醒类型"]
        # 呼吸暂停
        keyword_apnea1 = ["呼吸暂停+低通气"]
        keyword_apnea2 = ["所有.*暂停"]
        keyword_limboMovenments = ["睡眠期次数"]
        # 呼吸事件
        keyword_breathingEvent = ["AHI"]
        # 打鼾
        keyword_snoring = ["打鼾概要"]
        # 血氧
        keyword_oxygenSaturation = ["睡眠期平均血氧"]

        patterns = {
            tableType.Info: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_info]) + r')\b',
                re.IGNORECASE),
            tableType.FirstOrder: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_firstorder]) + r')\b',
                re.IGNORECASE),
            tableType.SleepStage: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_sleepStage]) + r')\b',
                re.IGNORECASE),
            tableType.Arousal: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_arousal]) + r')\b',
                re.IGNORECASE),
            tableType.Apnea1: re.compile(
                r'(?:{})'.format('|'.join(
                    [r'{}'.format(kw.replace('+', r'\+')) for kw in keyword_apnea1]
                )),
                re.IGNORECASE
            ),
            tableType.Apnea2: re.compile(
                r'(?:{})'.format('|'.join(keyword_apnea2)),
                re.IGNORECASE
            ),
            tableType.LimbMovements: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_limboMovenments]) + r')\b',
                re.IGNORECASE),
            tableType.BreathingEvent: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_breathingEvent]) + r')\b',
                re.IGNORECASE),
            tableType.Special: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_snoring]) + r')\b',
                re.IGNORECASE),
            tableType.Mixed: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_oxygenSaturation]) + r')\b',
                re.IGNORECASE),
        }
        table_type = tableType.Null
        for row in table:
            for cell in row:
                clean_cell = re.sub(r'\s+', '', str(cell))
                for table_type, pattern in patterns.items():
                    if pattern.search(clean_cell):
                        return table_type
        return tableType.Null

    def process_info_table(self,table, scan_mode=False):
        table_data = {}
        for row in table:
            """处理表格行（支持扫描模式和数据处理模式）"""
            cell_text = " | ".join(cell.strip() for cell in row)
            # 使用改进的正则表达式匹配键值对
            pattern = r'''
                ([^：]+?)        # 匹配键（非贪婪）
                \s*：\s*         # 匹配冒号及周围空格
                ((?:(?!\s*\||\s*$).)*)  # 匹配值（排除分隔符）
                (?=\s*\|?|\s*$)  # 前瞻断言
            '''

            for match in re.finditer(pattern, cell_text, re.X):
                raw_key, raw_value = match.groups()

                # 标准化键名：删除所有空格但保留符号
                clean_key = re.sub(r'\s+', '', raw_key)
                clean_key = re.sub(r'\|', '', clean_key)
                # 标准化值：保留原始内容仅清理空格
                clean_value = re.sub(r'\s*', '', raw_value).strip()
                if scan_mode:
                    # 扫描模式：注册字段
                    if clean_key not in table_data:
                        table_data[clean_key] = None
                else:
                    # 数据模式：存储值（最后出现的值会覆盖之前的）
                    table_data[clean_key] = clean_value
        return table_data

    def process_firstorder_table(self,table, scan_mode=False):
        table_data = {}
        for sublist in table:
            # 步长2遍历键值对
            for i in range(0, len(sublist), 2):
                key = sublist[i]
                value = sublist[i + 1]

                # 统一键名格式（可选）
                key = key.replace(" ", "").replace("（", "(").replace("）", ")").replace("\t",'')
                key = key.replace("总卧床时间TIB", "卧床时间(TIB)")
                key = key.replace("(次/分钟)","")

                # 转换数值类型（可选）
                table_data[key] = value

        return table_data


    def process_sleepstage_table(self,table, scan_mode = False):
        """处理睡眠分期表格"""
        # 创建结果字典
        result = {}

        # 处理每一行数据
        for row in table[1:]:
            # 清理分期名称
            stage = re.sub(r'\s+', '', row[0].strip())
            duration = float(row[1])
            percent = float(row[2])

            # 生成列名
            duration_col = f"{stage}持续时间(min)"
            percent_col = f"{stage}%睡眠时间(/TST)"

            # 存储数据
            result[duration_col] = duration
            result[percent_col] = percent
        return result

    def process_arousal_table(self,table, scan_mode=False):
        """处理微觉醒相关表格"""
        result = {}
        # 标准化表头
        headers = [re.sub(r'\(/TST\)', '', h).strip() for h in table[0]]

        # 确定数值列位置
        value_columns = []
        for idx, h in enumerate(headers):
            if h in ['REM', 'NREM', '次数', '指数']:
                value_columns.append((idx, h))

        # 处理数据行
        for row in table[1:]:
            arousal_type = row[0].strip()
            # 标准化类型名称
            arousal_type = re.sub(r'\s+', '', arousal_type)
            if arousal_type == 'Total':
                arousal_type = '微觉醒总数'

            for col_idx, col_name in value_columns:
                # 生成标准列名
                combined_col = f"{arousal_type}{col_name}"
                try:
                    value = float(row[col_idx])
                except (ValueError, IndexError):
                    value = 0.0

                result[combined_col] = value

        return result

    def process_apnea1_table(self,table, scan_mode=False):
        result = {}
        # 提取并标准化表头
        headers = [re.sub('\s*','',h).strip() for h in table[0]]

        # 处理数据行
        for row in table[1:]:
            apnea1_type = row[0].strip()
            apnea1_type = re.sub(r'\s+', '', apnea1_type)
            apnea1_type = re.sub(r'AHI\(/hr\)', r'AHI(/h)', apnea1_type)

            for col_idx in range(1, len(row)):
                # 生成标准化列名
                raw_col = headers[col_idx]
                std_col = f"{apnea1_type}{raw_col}"

                # 数值处理（新增空值和-处理）
                value = row[col_idx].strip()
                result[std_col] = value
        return result

    def process_apnea2_table(self,table, scan_mode=False):
        result = {}

        """处理呼吸事件表格并整合数据"""
        # 列名映射表（表二 -> 表一）
        column_mapping = {
            '阻塞性': '阻塞性呼吸暂停',
            '混合性': '混合性呼吸暂停',
            '中枢性': '中枢性呼吸暂停',
            '所有暂停': '所有呼吸暂停',
            '低通气': '所有低通气'
        }

        # 初始化结果容器
        result = {}
        key = '所有暂停'

        headers = [re.sub(r'\n', '', cell) for cell in table[0]]
        for row in table[1:]:
            param = row[0]
            for idx in range(1, len(row)):
                # 转换列名
                original_col = headers[idx]
                col_name = f"{headers[idx]}{param}"
                if key in table[0]:
                    mapped_col = column_mapping.get(original_col, original_col).strip()
                    col_name = f"{mapped_col}{param}"

                col_name = col_name.replace(" ", "").replace("（", "(").replace("）", ")")
                col_name = re.sub(r'\(sec\)', '(s)', col_name)  # 单位转换

                # 数值处理
                value = row[idx].strip()
                result[col_name] = '/' if value in ['-', ''] else value

        return result

    def process_limbomovements_table(self,table, scan_mode=False):
        result = {}
        column_mapping = {
            '睡眠期指数(/TST)': '睡眠期指数',
            '睡眠期指数': '睡眠期指数'
        }
        # 标准化表头
        headers = [column_mapping.get(col, col) for col in table[0]]

        # 处理数据行
        for row in table[1:]:
            event_type = row[0].strip()
            # 清理特殊字符并生成列名
            clean_type = event_type.replace("相关", "_").replace(" ", "")

            # 动态生成列名
            count_col = f"{clean_type}_次数"
            index_col = f"{clean_type}_指数"

            # 存储数据（后续表格覆盖前期数据）
            result[count_col] = row[1] if row[1] not in ['', '-'] else '/'
            result[index_col] = row[2] if row[2] not in ['', '-'] else '/'
        return result

    def process_breathingevent_table(self,table, scan_mode=False):
        result = {}
        return result

    def process_snoring_table(self,table, scan_mode=False):
        result = {}
        return result

    def process_oxygenSaturation_table(self,table, scan_mode=False):
        result = {}
        return result

    def process_table_data(self,table, table_type):
        if table_type == tableType.Info:
            return self.process_info_table(table)
        elif table_type == tableType.FirstOrder:
            return self.process_firstorder_table(table)
        elif table_type == tableType.SleepStage:
            return self.process_sleepstage_table(table)
        elif table_type == tableType.Arousal:
            return self.process_arousal_table(table)
        elif table_type == tableType.Apnea1:
            return self.process_apnea1_table(table)
        elif table_type == tableType.Apnea2:
            return self.process_apnea2_table(table)
        elif table_type == tableType.LimbMovements:
            return self.process_limbomovements_table(table)
        elif table_type == tableType.BreathingEvent:
            return self.process_breathingevent_table(table)
        elif table_type == tableType.Snoring:
            return self.process_snoring_table(table)
        elif table_type == tableType.OxygenSaturation:
            return self.process_oxygenSaturation_table(table)
        else:
            return {}

    def load_config(self,yaml_path):
        """加载YAML配置文件"""
        with open(yaml_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
        return ['文件名'] + config['Param']  # 第一列为文件名


    def rtf_to_docx(self,rtf_path):
        """转换RTF为DOCX"""
        temp_dir = os.path.dirname(rtf_path)
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'docx',
            '--outdir', temp_dir, rtf_path
        ], check=True, capture_output=True)
        return os.path.splitext(rtf_path)[0] + ".docx"

    def iter_block_items(self,parent):
        """
        生成父元素中的每个段落和表格元素。
        """
        for child in parent:
            if isinstance(child, OxmlElement):
                if child.tag.endswith('p'):
                    yield child
                elif child.tag.endswith('tbl'):
                    yield child

    # 将doc中的元素存入element中，准备按顺序处理
    def print_docx_content(self,doc):
        # 构建元素序列
        elements = []
        for element in doc.element.body.iterchildren():
            if isinstance(element, CT_P):
                # 通过元素创建段落对象
                para = Paragraph(element, doc.part)
                elements.append(('paragraph', para))
            elif isinstance(element, CT_Tbl):
                # 通过元素创建表格对象
                tbl = Table(element, doc.part)
                elements.append(('table', tbl))

        para_number = 0
        table_number = 0
        # 按原始顺序处理元素
        for elem_type, elem_obj in elements:
            if elem_type == 'paragraph':
                text = elem_obj.text.strip()
                para_number +=1
                if text:
                    print(f"[段落] {para_number}: {text}")
            elif elem_type == 'table':
                table_number += 1
                print(f"\n[表格] {table_number}")
                for row in elem_obj.rows:
                    row_data = [
                        " ".join(run.text for run in cell.paragraphs[0].runs)
                        for cell in row.cells
                    ]
                    print(" | ".join(row_data))
                print("-" * 40)



    def extract_docx_data(self,docx_path, fields):
        """从DOCX提取目标数据"""
        doc = Document(docx_path)
        data = {field: "" for field in fields}

        # print_docx_content(doc)

        # 提取所有段落文本
        full_text = []
        current_section = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith("#"):
                current_section = text[1:].strip()
            else:
                full_text.append((current_section, text))

        # 提取所有表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text.strip() for cell in row.cells])
            tables.append(table_data)

        for table in tables:
            table_type = self.judge_table_type(table)
            table_data = self.process_table_data(table,table_type)
            for field in fields:
                for key_data in table_data:
                    if field == key_data:
                        data[field] = table_data[key_data]
            if table_type == tableType.LimbMovements:
                print(table_type,":\ntable_data:\n",table_data,"\ntable:\n",table)

        #print(data)
        return data


    def process_files(self,folder_path):
        """处理文件夹中的所有RTF文件"""
        # 初始化Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "合并数据"

        # 获取字段配置
        config_path = os.path.join(os.getcwd(), YAML_CONFIG)
        fields = self.load_config(config_path)

        # 创建表头
        for col_idx, field in enumerate(fields, 1):
            ws.cell(row=1, column=col_idx, value=field)

        # 处理文件
        row_idx = 2
        for filename in os.listdir(folder_path):
            if not filename.lower().endswith('.rtf'):
                continue

            print(f"正在处理 {filename}......")

            filepath = os.path.join(folder_path, filename)
            try:
                # 转换文件格式
                docx_path = self.rtf_to_docx(filepath)

                # 提取数据
                file_data = self.extract_docx_data(docx_path, fields)
                file_data['文件名'] = os.path.splitext(filename)[0]

                # 写入Excel
                for col_idx, field in enumerate(fields, 1):
                    ws.cell(row=row_idx, column=col_idx, value=file_data.get(field, ""))

                row_idx += 1
                os.remove(docx_path)  # 清理临时文件
                print(f"文件{filename}处理结束")

            except Exception as e:
                print(f"处理失败 {filename}: {str(e)}")
                continue

        # 自动调整列宽
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    cell_value = str(cell.value)
                    if len(cell_value) > max_length:
                        max_length = len(cell_value)
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column].width = adjusted_width

        excel_name = os.path.basename(folder_path)
        excel_name = excel_name + ".xlsx"
        print(excel_name)
        excel_output = os.path.join(folder_path, excel_name)
        print(excel_output)
        wb.save(excel_output)
        print(f"处理完成！结果已保存至{excel_output}")


if __name__ == "__main__":
    # 使用示例
    rtf_parser = RtfDataParser()
    folder_path =  r"D:\workshop\unittest\testdata"
    rtf_parser.process_files(folder_path)
    print(f"处理完成！结果已保存")

