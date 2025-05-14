import os
import yaml
import subprocess
from docx import Document
from openpyxl import Workbook
from docx.oxml import OxmlElement
import enum
import re
from collections import OrderedDict

from log_processor import LogManager
import time
import threading
from queue import Queue

# 配置文件路径
YAML_CONFIG = "MedicalReportParameters.yml"
class tableType(enum.Enum):
    Null = enum.auto()
    Info = enum.auto()
    FirstOrder = enum.auto()
    SleepStage = enum.auto()
    Arousal = enum.auto()
    Apnea1 = enum.auto()
    Apnea2 = enum.auto()
    LimbMovements = enum.auto()
    BreathingEvent = enum.auto()
    OxygenSaturation = enum.auto()
    Snoring = enum.auto()

def convert_time(time_str):
    try:
        if ":" in time_str:  # 处理类似"0:12:2.0"的格式
            parts = list(map(float, time_str.split(":")))
            return round(parts[0] * 60 + parts[1] + parts[2] / 60, 2)
        return float(time_str)
    except:
        return None

def extract_number(value):
    try:
        # 移除可能存在的百分号
        cleaned = str(value).replace('%', '').strip()
        return float(cleaned)
    except:
        return None
class RTFParser:
    def __init__(self,log_queue,stop_event):
        self.logger = LogManager().get_logger()
        self.log_queue = log_queue
        self._stop_event = stop_event
        self._stop_event = threading.Event()

    def judge_table_type(self,table):
        keyword_info = ["姓名"]
        keyword_firstorder = ["熄灯时间", "睡眠期平均心率"]
        # 睡眠分期
        keyword_sleepStage = ["睡眠时间"]
        # 微觉醒类型
        keyword_arousal = ["微觉醒类型"]
        # 呼吸暂停
        keyword_apnea1 = ["呼吸暂停+低通气"]
        keyword_apnea2 = ["所有.*暂停"]
        keyword_limboMovenments = ["睡眠期次数"]
        # 呼吸事件
        keyword_breathingEvent = ["AHI"]
        # 打鼾
        keyword_snoring = ["打鼾概要"]
        # 血氧
        keyword_oxygenSaturation = ["睡眠期平均血氧"]

        patterns = {
            tableType.Info: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_info]) + r')\b',
                re.IGNORECASE),
            tableType.FirstOrder: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_firstorder]) + r')\b',
                re.IGNORECASE),
            tableType.SleepStage: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_sleepStage]) + r')\b',
                re.IGNORECASE),
            tableType.Arousal: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_arousal]) + r')\b',
                re.IGNORECASE),
            tableType.Apnea1: re.compile(
                r'(?:{})'.format('|'.join(
                    [r'{}'.format(kw.replace('+', r'\+')) for kw in keyword_apnea1]
                )),
                re.IGNORECASE
            ),
            tableType.Apnea2: re.compile(
                r'(?:{})'.format('|'.join(keyword_apnea2)),
                re.IGNORECASE
            ),
            tableType.LimbMovements: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_limboMovenments]) + r')\b',
                re.IGNORECASE),
            tableType.BreathingEvent: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_breathingEvent]) + r')\b',
                re.IGNORECASE),
            tableType.Snoring: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_snoring]) + r')\b',
                re.IGNORECASE),
            tableType.OxygenSaturation: re.compile(
                r'\b(?:' + '|'.join([re.escape(kw).replace('\\', '\\\\') for kw in keyword_oxygenSaturation]) + r')\b',
                re.IGNORECASE),
        }
        table_type = tableType.Null
        for row in table:
            for cell in row:
                clean_cell = re.sub(r'\s+', '', str(cell))
                for table_type, pattern in patterns.items():
                    if pattern.search(clean_cell):
                        return table_type
        return tableType.Null

    def process_info_table(self,table, scan_mode=False):
        table_data = {}
        for row in table:
            """处理表格行（支持扫描模式和数据处理模式）"""
            cell_text = " | ".join(cell.strip() for cell in row)
            # 使用改进的正则表达式匹配键值对
            pattern = r'''
                ([^：]+?)        # 匹配键（非贪婪）
                \s*：\s*         # 匹配冒号及周围空格
                ((?:(?!\s*\||\s*$).)*)  # 匹配值（排除分隔符）
                (?=\s*\|?|\s*$)  # 前瞻断言
            '''

            for match in re.finditer(pattern, cell_text, re.X):
                raw_key, raw_value = match.groups()

                # 标准化键名：删除所有空格但保留符号
                clean_key = re.sub(r'\s+', '', raw_key)
                clean_key = re.sub(r'\|', '', clean_key)
                # 标准化值：保留原始内容仅清理空格
                clean_value = re.sub(r'\s*', '', raw_value).strip()
                if scan_mode:
                    # 扫描模式：注册字段
                    if clean_key not in table_data:
                        table_data[clean_key] = None
                else:
                    # 数据模式：存储值（最后出现的值会覆盖之前的）
                    table_data[clean_key] = clean_value
        return table_data

    def process_firstorder_table(self,table, scan_mode=False):
        table_data = {}
        for sublist in table:
            # 步长2遍历键值对
            for i in range(0, len(sublist), 2):
                key = sublist[i]
                value = sublist[i + 1]

                # 统一键名格式（可选）
                key = key.replace(" ", "").replace("（", "(").replace("）", ")").replace("\t",'')
                key = key.replace("总卧床时间TIB", "卧床时间(TIB)")
                key = key.replace("(次/分钟)","")

                # 转换数值类型（可选）
                table_data[key] = value

        return table_data


    def process_sleepstage_table(self,table, scan_mode = False):
        """处理睡眠分期表格"""
        # 创建结果字典
        result = {}

        # 处理每一行数据
        for row in table[1:]:
            # 清理分期名称
            stage = re.sub(r'\s+', '', row[0].strip())
            duration = float(row[1])
            percent = float(row[2])

            # 生成列名
            duration_col = f"{stage}持续时间(min)"
            percent_col = f"{stage}%睡眠时间(/TST)"

            # 存储数据
            result[duration_col] = duration
            result[percent_col] = percent
        return result

    def process_arousal_table(self,table, scan_mode=False):
        """处理微觉醒相关表格"""
        result = {}
        # 标准化表头
        headers = [re.sub(r'\(/TST\)', '', h).strip() for h in table[0]]

        # 确定数值列位置
        value_columns = []
        for idx, h in enumerate(headers):
            if h in ['REM', 'NREM', '次数', '指数']:
                value_columns.append((idx, h))

        # 处理数据行
        for row in table[1:]:
            arousal_type = row[0].strip()
            # 标准化类型名称
            arousal_type = re.sub(r'\s+', '', arousal_type)
            if arousal_type == 'Total':
                arousal_type = '微觉醒总数'

            for col_idx, col_name in value_columns:
                # 生成标准列名
                combined_col = f"{arousal_type}{col_name}"
                try:
                    value = float(row[col_idx])
                except (ValueError, IndexError):
                    value = 0.0

                result[combined_col] = value

        return result

    def process_apnea1_table(self,table, scan_mode=False):
        result = {}
        # 提取并标准化表头
        headers = [re.sub('\s*','',h).strip() for h in table[0]]

        # 处理数据行
        for row in table[1:]:
            apnea1_type = row[0].strip()
            apnea1_type = re.sub(r'\s+', '', apnea1_type)
            apnea1_type = re.sub(r'AHI\(/hr\)', r'AHI(/h)', apnea1_type)

            for col_idx in range(1, len(row)):
                # 生成标准化列名
                raw_col = headers[col_idx]
                std_col = f"{apnea1_type}{raw_col}"

                # 数值处理（新增空值和-处理）
                value = row[col_idx].strip()
                result[std_col] = value
        return result

    def process_apnea2_table(self,table, scan_mode=False):
        result = {}

        """处理呼吸事件表格并整合数据"""
        # 列名映射表（表二 -> 表一）
        column_mapping = {
            '阻塞性': '阻塞性呼吸暂停',
            '混合性': '混合性呼吸暂停',
            '中枢性': '中枢性呼吸暂停',
            '所有暂停': '所有呼吸暂停',
            '低通气': '所有低通气'
        }

        # 初始化结果容器
        result = {}
        key = '所有暂停'

        headers = [re.sub(r'\n', '', cell) for cell in table[0]]
        for row in table[1:]:
            param = row[0]
            for idx in range(1, len(row)):
                # 转换列名
                original_col = headers[idx]
                col_name = f"{headers[idx]}{param}"
                if key in table[0]:
                    mapped_col = column_mapping.get(original_col, original_col).strip()
                    col_name = f"{mapped_col}{param}"

                col_name = col_name.replace(" ", "").replace("（", "(").replace("）", ")")
                col_name = re.sub(r'\(sec\)', '(s)', col_name)  # 单位转换

                # 数值处理
                value = row[idx].strip()
                result[col_name] = '/' if value in ['-', ''] else value

        return result

    def process_limbomovements_table(self,table, scan_mode=False):
        result = {}

        # 列名标准化映射
        column_mapping = {
            "睡眠期指数(/TST)": "睡眠期指数(/TST)",
            "睡眠期指数": "睡眠期指数(/TST)"
        }

        # 结果容器（保持列顺序）

        # 标准化表头
        headers = [column_mapping.get(col.strip(), col.strip()) for col in table[0]]

        # 处理数据行
        for row in table[1:]:
            # 清理类型名称
            event_type = re.sub(r'\s+', '', row[0].strip())
            # 动态生成列名
            count_col = f"{event_type}睡眠期次数"
            index_col = f"{event_type}睡眠期指数(/TST)"
            # 数据清洗（空值和-转为/）
            count_value = row[1] if row[1] not in ['', '-'] else '/'
            index_value = row[2] if row[2] not in ['', '-'] else '/'

            # 存储数据
            result[count_col] = count_value
            result[index_col] = index_value

        return result

    def process_breathingevent_table(self,table, scan_mode=False):
        """处理呼吸事件表格（完整保留所有体位数据）"""
        # 预定义标准体位和指标
        POSITIONS = ['俯卧', '左侧', '右侧', '仰卧']
        METRICS = [
            '阻塞性呼吸暂停', '混合性呼吸暂停', '中枢性呼吸暂停',
            '低通气', 'AHI', '睡眠时间%', '持续时间(min)'
        ]

        # 初始化结果容器（动态扩展列）
        result = OrderedDict()

        # === 列名映射规则 ===
        column_mapping = {
            r'阻塞性[\s（]*': '阻塞性呼吸暂停',
            r'混合性[\s（]*': '混合性呼吸暂停',
            r'中枢性[\s（]*': '中枢性呼吸暂停',
            r'低通气[\s（]*': '低通气',
            r'^AHI$': 'AHI',
            r'睡眠时间': '睡眠时间%',
            r'持续时间[\s（]*min': '持续时间(min)'
        }

        # === 表头处理 ===
        header_map = {}
        for orig_col in table[0]:
            # 深度清洗列名
            cleaned_col = re.sub(r'[\n（）()]', '', orig_col).strip()
            # 动态匹配列名规则
            for pattern, mapped in column_mapping.items():
                if re.search(pattern, cleaned_col):
                    header_map[orig_col] = mapped
                    break
        # === 数据处理 ===
        for row in table[1:]:
            position = re.sub(r'\s+', '', row[0])
            if position not in POSITIONS:
                continue

            for idx in range(1, len(row)):
                orig_col = table[0][idx]
                metric = header_map.get(orig_col)

                if metric:
                    # 生成完整列名
                    col_name = f"{position}{metric}"
                    raw_value = row[idx].strip()

                    # 数据清洗
                    if raw_value in ['', '-', 'NA']:
                        final_value = '/'
                    else:
                        try:
                            final_value = float(raw_value) if '.' in raw_value else int(raw_value)
                        except:
                            final_value = '/'

                    # 直接存储所有值（包括0）
                    result[col_name] = final_value

        # === 补全所有可能的列 ===
        full_columns = [f"{pos}{metric}" for pos in POSITIONS for metric in METRICS]
        for col in full_columns:
            if col not in result:
                result[col] = '/'

        # 保持列顺序
        result = OrderedDict((col, result.get(col, '/')) for col in full_columns)
        return result


    def process_snoring_table(self,table, scan_mode=False):
        result = {}
        for i in range(0, len(table[1]), 2):
            key = table[1][i]
            clean_key = re.sub(r'（睡眠期）', '', key)
            value = table[1][i + 1]
            result[clean_key] = value
        return result


    def process_oxygenSaturation_table(self,table, scan_mode=False):
        result = {}

        for row in table:
            # 处理前四个参数
            for i, cell in enumerate(row):
                if "睡眠期平均血氧 (%)" in cell:
                    result["睡眠期平均血氧"] = extract_number(row[i + 1]) if i + 1 < len(row) else None
                if "清醒期平均SpO2 (%)" in cell:
                    result["清醒期平均SpO2(%)"] = extract_number(row[i + 1]) if i + 1 < len(row) else None
                if "睡眠期最低血氧 (%)" in cell:
                    result["睡眠期最低血氧(%)"] = extract_number(row[i + 1]) if i + 1 < len(row) else None
                if "氧减" in cell and "指数" in cell:
                    result["氧减＞3%指数(/h)(ODI)"] = extract_number(row[i + 1]) if i + 1 < len(row) else None

            # 处理血氧饱和度水平数据（优化定位逻辑）
            if row[0].startswith('低于'):
                key_type = row[0].split(' ')[0]  # 如"低于95%"

                # 统一提取规则：时间取第2列，占比取最后一列
                if "时间（min）" in row[0]:
                    # 时间值处理
                    time_val = row[1] if len(row) > 1 else None
                    # 占比值处理
                    percent_val = row[-1] if len(row) > 1 else None

                    # 根据百分比级别存储数据
                    if "95%" in key_type:
                        result["血氧饱和度水平低于95%时间(min)"] = convert_time(time_val)
                        result["血氧饱和度水平低于95%时间占比(%)"] = extract_number(percent_val)
                    elif "90%" in key_type:
                        result["血氧饱和度水平低于90%时间(min)"] = convert_time(time_val)
                        result["血氧饱和度水平低于90%时间占比(%)"] = extract_number(percent_val)
                    elif "85%" in key_type:
                        result["血氧饱和度水平低于85%时间(min)"] = convert_time(time_val)
                        result["血氧饱和度水平低于85%时间占比(%)"] = extract_number(percent_val)
                    elif "80%" in key_type:
                        result["血氧饱和度水平低于80%时间(min)"] = convert_time(time_val)
                        result["血氧饱和度水平低于80%时间占比(%)"] = extract_number(percent_val)

        return result

    def process_table_data(self,table, table_type):
        if table_type == tableType.Info:
            return self.process_info_table(table)
        elif table_type == tableType.FirstOrder:
            return self.process_firstorder_table(table)
        elif table_type == tableType.SleepStage:
            return self.process_sleepstage_table(table)
        elif table_type == tableType.Arousal:
            return self.process_arousal_table(table)
        elif table_type == tableType.Apnea1:
            return self.process_apnea1_table(table)
        elif table_type == tableType.Apnea2:
            return self.process_apnea2_table(table)
        elif table_type == tableType.LimbMovements:
            return self.process_limbomovements_table(table)
        elif table_type == tableType.BreathingEvent:
            return self.process_breathingevent_table(table)
        elif table_type == tableType.Snoring:
            return self.process_snoring_table(table)
        elif table_type == tableType.OxygenSaturation:
            return self.process_oxygenSaturation_table(table)
        else:
            return {}

    def load_config(self,yaml_path):
        """加载YAML配置文件"""
        with open(yaml_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
        return ['文件名'] + config['Param']  # 第一列为文件名


    def rtf_to_docx(self,rtf_path):
        """转换RTF为DOCX"""
        temp_dir = os.path.dirname(rtf_path)
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'docx',
            '--outdir', temp_dir, rtf_path
        ], check=True, capture_output=True)
        return os.path.splitext(rtf_path)[0] + ".docx"

    def iter_block_items(self,parent):
        """
        生成父元素中的每个段落和表格元素。
        """
        for child in parent:
            if isinstance(child, OxmlElement):
                if child.tag.endswith('p'):
                    yield child
                elif child.tag.endswith('tbl'):
                    yield child

    def extract_data(self,paragraphs):
        """从段落数据中提取目标字段"""
        data = {
            "AHI(次/h)": None,
            "OAHI(次/h)": None,
            "OAI(次/h)": None,
            "睡眠期间血氧＜90%的累计时间(min)": None,
            "睡眠期间血氧＜90%的累计时间占比": None,
            "结论": [],
            "诊断": []
        }

        # 状态标志
        in_conclusion = False
        in_diagnosis = False

        for _, text in paragraphs:
            text = text.strip()
            if not text:
                in_conclusion = False
                in_diagnosis = False
                continue

            # 提取数值型数据
            if "AHI" in text:
                if match := re.search(r"AHI.*?=([\d.]+)", text):
                    data["AHI(次/h)"] = float(match.group(1))
            if match := re.search(r"OAHI.*?=([\d.]+)", text):
                data["OAHI(次/h)"] = float(match.group(1))
            if match := re.search(r"OAI.*?=([\d.]+)", text):
                data["OAI(次/h)"] = float(match.group(1))

            # 血氧数据提取
            if "血氧<90%" in text or "血氧＜90%" in text:
                parts = re.split(r"[；;]", text)
                for part in parts:
                    if "时间" in part and "min" not in part:  # 处理第三个文档的特殊格式
                        if match := re.search(r"([\d:\.]+)", part):
                            data["睡眠期间血氧＜90%的累计时间(min)"] = convert_time(match.group(1))
                    elif "时间" in part:
                        if match := re.search(r"([\d.]+)\s*min", part):
                            data["睡眠期间血氧＜90%的累计时间(min)"] = float(match.group(1))
                    if "占比" in part:
                        if match := re.search(r"([\d.]+)%?", part):
                            data["睡眠期间血氧＜90%的累计时间占比"] = float(match.group(1))

            # 结论和诊断处理
            if text.startswith("结论："):
                in_conclusion = True
                in_diagnosis = False
                data["结论"].append(text.replace("结论：", "").strip())
                continue
            if text.startswith("诊断："):
                in_diagnosis = True
                in_conclusion = False
                data["诊断"].append(text.replace("诊断：", "").strip())
                continue

            if in_conclusion:
                data["结论"].append(text)
            if in_diagnosis:
                data["诊断"].append(text)

        # 合并文本字段
        data["结论"] = " ".join(data["结论"]) if data["结论"] else None
        data["诊断"] = " ".join(data["诊断"]) if data["诊断"] else None

        return data


    def extract_docx_data(self,docx_path, fields):
        """从DOCX提取目标数据"""
        doc = Document(docx_path)
        data = {field: "" for field in fields}

        # print_docx_content(doc)

        # 提取所有段落文本
        full_text = []
        current_section = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith("#"):
                current_section = text[1:].strip()
            else:
                full_text.append((current_section, text))

        doc_data = self.extract_data(full_text)
        for field in fields:
            for key_data in doc_data:
                if field == key_data:
                    data[field] = doc_data[key_data]

        # 提取所有表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text.strip() for cell in row.cells])
            tables.append(table_data)

        for table in tables:
            table_type = self.judge_table_type(table)
            table_data = self.process_table_data(table,table_type)
            for field in fields:
                for key_data in table_data:
                    if field == key_data:
                        data[field] = table_data[key_data]
            if not table_type == table_type.Null:
                debug_msg = f"{table_type},:\ntable_data:\n{table_data}\ntable:\n{table}"
                self.logger.debug("%s", debug_msg)

        self.logger.debug(data)
        return data

    def stop(self):
        """停止解析"""
        self._stop_event.set()

    def process_files(self,folder_path):
        """处理文件夹中的所有RTF文件"""
        # 初始化Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "合并数据"

        # 获取字段配置
        config_path = os.path.join(os.getcwd(), YAML_CONFIG)
        fields = self.load_config(config_path)

        # 创建表头
        for col_idx, field in enumerate(fields, 1):
            ws.cell(row=1, column=col_idx, value=field)

        # 处理文件
        row_idx = 2
        for filename in os.listdir(folder_path):
            if not filename.lower().endswith('.rtf'):
                continue

            self.logger.info(f"正在处理 {filename}......")
            if self._stop_event.is_set():
                self.logger.info("接受到停止请求，任务已经终止")
                return False

            filepath = os.path.join(folder_path, filename)
            try:
                # 转换文件格式
                docx_path = self.rtf_to_docx(filepath)

                # 提取数据
                file_data = self.extract_docx_data(docx_path, fields)
                file_data['文件名'] = os.path.splitext(filename)[0]

                # 写入Excel
                for col_idx, field in enumerate(fields, 1):
                    ws.cell(row=row_idx, column=col_idx, value=file_data.get(field, ""))

                row_idx += 1
                os.remove(docx_path)  # 清理临时文件
                self.logger.info(f"文件{filename}处理结束")

            except Exception as e:
                self.logger.error(f"处理失败 {filename}: {str(e)}")
                continue

        # 自动调整列宽
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    cell_value = str(cell.value)
                    if len(cell_value) > max_length:
                        max_length = len(cell_value)
                except:
                    pass
            adjusted_width = (max_length + 2) * 1.2
            ws.column_dimensions[column].width = adjusted_width

        excel_name = os.path.basename(folder_path)
        excel_name = excel_name + ".xlsx"
        excel_output = os.path.join(folder_path, excel_name)
        wb.save(excel_output)
        self.logger.info(f"处理完成！结果已保存至{excel_output}")


if __name__ == "__main__":
    # 使用示例
    rtf_parser = RTFParser()
    folder_path =  r"D:\workshop\数据测试用PSG data\数据测试用PSG data"
    rtf_parser.process_files(folder_path)
    print(f"处理完成！结果已保存")